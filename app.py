from flask import Flask, render_template, request
import PyPDF2
import docx
import os
import re

app = Flask(__name__)

# -------------------------------
# Upload folder setup
# -------------------------------
UPLOAD_FOLDER = "uploads"
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# -------------------------------
# Keywords with weights
# -------------------------------
keywords = {
    "python": 3,
    "machine learning": 4,
    "data analysis": 3,
    "sql": 2,
    "communication": 1,
    "teamwork": 1
}

# -------------------------------
# Extract text from resume
# -------------------------------
def extract_text(filepath):
    text = ""

    try:
        if filepath.endswith('.pdf'):
            with open(filepath, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                for page in pdf.pages:
                    if page.extract_text():
                        text += page.extract_text() + " "

        elif filepath.endswith('.docx'):
            doc = docx.Document(filepath)

            # Extract paragraphs
            for para in doc.paragraphs:
                text += para.text + " "

            # 🔥 Extract tables (VERY IMPORTANT)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "

    except Exception as e:
        print("Error reading file:", e)

    return text.lower()


# -------------------------------
# Clean text
# -------------------------------
def clean_text(text):
    return re.sub(r'[^a-zA-Z0-9\s]', '', text)


# -------------------------------
# Score resume (0–100 guaranteed)
# -------------------------------
def score_resume(text):
    total_weight = sum(keywords.values())
    matched_weight = 0
    matched = []

    for word, weight in keywords.items():
        if word in text:
            matched_weight += weight
            matched.append(word)

    percentage = (matched_weight / total_weight) * 100 if total_weight != 0 else 0

    return round(percentage, 2), matched


# -------------------------------
# Routes
# -------------------------------
@app.route('/')
def home():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload():
    files = request.files.getlist('resumes')

    results = []

    for file in files:
        if file.filename == "":
            continue

        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)

        text = extract_text(filepath)
        text = clean_text(text)

        score, matched = score_resume(text)

        # 🔥 Selection logic
        if score >= 70:
            status = "HIGHLY SHORTLISTED"
        elif score >= 50:
            status = "SHORTLISTED"
        else:
            status = "REJECTED"

        results.append({
            "name": file.filename,
            "score": score,
            "matched": matched if matched else ["No relevant skills found"],
            "status": status
        })

    # Sort by score (ranking)
    results = sorted(results, key=lambda x: x['score'], reverse=True)

    return render_template('result.html', results=results)


# -------------------------------
# Run app
# -------------------------------
if __name__ == '__main__':
    app.run(debug=True)
    